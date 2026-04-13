"""
Convert HLD.md and ADD.md to formatted Word documents.
Run from the repo root:
    python generate_hld_add_docs.py
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_col_width(table, col_idx: int, width_cm: float):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def style_run(run, bold=False, italic=False, code=False,
              size_pt=None, color_hex=None):
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
    if size_pt:
        run.font.size = Pt(size_pt)
    if color_hex:
        r, g, b = (int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        run.font.color.rgb = RGBColor(r, g, b)


def add_inline_formatted(para, text: str):
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
        else:
            para.add_run(part)


# ── document setup ────────────────────────────────────────────────────────────

def setup_styles(doc: Document, accent_rgb=(0x1F, 0x35, 0x64),
                 h2_rgb=(0x2E, 0x74, 0xB5)):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(*accent_rgb)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(*h2_rgb)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x1F, 0x60, 0x9A)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(3)

    h4 = styles["Heading 4"]
    h4.font.name = "Calibri"
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(2)


def add_cover_page(doc: Document, title: str, subtitle: str,
                   meta_rows: list, header_hex="2E74B5", light_hex="EBF3FB",
                   title_rgb=(0x1F, 0x35, 0x64), sub_rgb=(0x2E, 0x74, 0xB5)):
    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(*title_rgb)
    run.font.name = "Calibri"

    # Subtitle
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(*sub_rgb)
    run2.font.name = "Calibri"

    doc.add_paragraph()
    doc.add_paragraph()

    # Metadata table
    meta = doc.add_table(rows=len(meta_rows), cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    for i, (label, value) in enumerate(meta_rows):
        meta.rows[i].cells[0].text = label
        meta.rows[i].cells[1].text = value
        set_cell_bg(meta.rows[i].cells[0], header_hex)
        meta.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        meta.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(meta.rows[i].cells[1], light_hex)

    doc.add_page_break()


# ── markdown → docx parser ───────────────────────────────────────────────────

def parse_markdown(doc: Document, md_text: str, header_hex="2E74B5",
                   light_hex="EBF3FB"):
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines = []

    def flush_code():
        nonlocal code_lines
        if not code_lines:
            return
        combined = "\n".join(code_lines)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)
        run = p.add_run(combined)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
        code_lines.clear()

    while i < len(lines):
        line = lines[i]

        # ── code block toggle ──
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
            else:
                flush_code()
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── headings ──
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)

        # ── horizontal rule ──
        elif line.strip() == "---":
            add_horizontal_rule(doc)

        # ── table ──
        elif line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip().strip("|").split("|")]
                rows.append(cells)

            data_rows = [r for r in rows if not all(re.match(r"^[-: ]+$", c) for c in r)]
            if not data_rows:
                continue

            num_cols = max(len(r) for r in data_rows)
            for r in data_rows:
                while len(r) < num_cols:
                    r.append("")

            tbl = doc.add_table(rows=len(data_rows), cols=num_cols)
            tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

            for ri, row in enumerate(data_rows):
                for ci, cell_text in enumerate(row):
                    cell = tbl.rows[ri].cells[ci]
                    cell.paragraphs[0].clear()
                    p = cell.paragraphs[0]
                    if ri == 0:
                        set_cell_bg(cell, header_hex)
                        run = p.add_run(cell_text)
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(9.5)
                        run.font.name = "Calibri"
                    else:
                        set_cell_bg(cell, light_hex if ri % 2 == 0 else "FFFFFF")
                        add_inline_formatted(p, cell_text)
                        for run in p.runs:
                            run.font.size = Pt(9.5)
                            run.font.name = "Calibri"

            doc.add_paragraph()
            continue

        # ── bullet / numbered list ──
        elif re.match(r"^(\s*)[-*+] ", line):
            indent = len(line) - len(line.lstrip())
            text = re.sub(r"^\s*[-*+] ", "", line)
            style = "List Bullet" if indent == 0 else "List Bullet 2"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.3)
            add_inline_formatted(p, text)

        elif re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line)
            p = doc.add_paragraph(style="List Number")
            add_inline_formatted(p, text)

        # ── blank line ──
        elif line.strip() == "":
            pass

        # ── normal paragraph ──
        else:
            p = doc.add_paragraph()
            add_inline_formatted(p, line)

        i += 1

    if in_code:
        flush_code()


# ── builders ─────────────────────────────────────────────────────────────────

def build_doc(md_path: Path, out_path: Path, cover_title: str,
              cover_subtitle: str, meta_rows: list,
              header_hex="2E74B5", light_hex="EBF3FB"):
    print(f"Reading : {md_path}")
    md_text = md_path.read_text(encoding="utf-8")

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    setup_styles(doc)
    add_cover_page(doc, cover_title, cover_subtitle, meta_rows,
                   header_hex=header_hex, light_hex=light_hex)
    parse_markdown(doc, md_text, header_hex=header_hex, light_hex=light_hex)

    doc.save(str(out_path))
    print(f"Saved   : {out_path}")


def build_hld():
    docs_dir = Path(__file__).parent / "docs"
    build_doc(
        md_path       = docs_dir / "HLD.md",
        out_path      = docs_dir / "HLD_Scrum_Master_Digital_Worker.docx",
        cover_title   = "HIGH LEVEL DESIGN DOCUMENT",
        cover_subtitle= "Scrum Master Digital Worker\nSprint Status Agent",
        meta_rows     = [
            ("Version",   "1.0"),
            ("Date",      "2026-04-13"),
            ("Status",    "Approved"),
            ("Audience",  "Stakeholders, Product Owners, Tech Leads"),
            ("System",    "Jira + LangGraph + MCP Server"),
            ("Delivery",  "Slack / Confluence"),
        ],
        header_hex    = "2E74B5",
        light_hex     = "EBF3FB",
    )


def build_add():
    docs_dir = Path(__file__).parent / "docs"
    build_doc(
        md_path       = docs_dir / "ADD.md",
        out_path      = docs_dir / "ADD_Scrum_Master_Digital_Worker.docx",
        cover_title   = "ARCHITECTURE DESIGN DOCUMENT",
        cover_subtitle= "Scrum Master Digital Worker\nSprint Status Agent",
        meta_rows     = [
            ("Version",   "1.0"),
            ("Date",      "2026-04-13"),
            ("Status",    "Approved"),
            ("Audience",  "Senior Engineers, Architects"),
            ("System",    "Jira + LangGraph + MCP Server"),
            ("Delivery",  "Slack / Confluence"),
        ],
        header_hex    = "1F3564",
        light_hex     = "EEF1F8",
    )


# ── main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_hld()
    build_add()
    print("\nDone. Both Word documents written to docs/")
